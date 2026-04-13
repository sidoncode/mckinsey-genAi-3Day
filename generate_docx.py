"""Generate a comprehensive DOCX document for the 3-Day GenAI Training Program."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading_elm.append(shading)


def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return heading


def add_table_row(table, cells_text, bold=False, header=False, shade_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        para = cell.paragraphs[0]
        run = para.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "003366")
        elif shade_color:
            set_cell_shading(cell, shade_color)
    return row


def create_document():
    doc = Document()

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==================== TITLE PAGE ====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("3-Day Generative AI Training Program")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Comprehensive Course Guide\nTopics, Flow, Exercises & Tasks")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("McKinsey & Company | AI/ML Training Division")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Program Overview & 3-Day Flow",
        "2. Day 1: Foundations, Prompting & Evaluation",
        "   2.1 Session 1: Modern LLM Foundations",
        "   2.2 Session 2: Prompt Engineering for Agentic Behaviors",
        "   2.3 Session 3: Model Evaluation and Comparison",
        "   2.4 Session 4: Lab Review & Integration",
        "3. Day 2: LangChain, LangGraph & Multi-Agent Systems",
        "   3.1 Session 1: OpenAI API & Structured Outputs",
        "   3.2 Session 2: LangChain & Tool Integration",
        "   3.3 Session 3: LangGraph Orchestration",
        "   3.4 Session 4: Multi-Agent Workflows",
        "4. Day 3: RAG, Deployment & Capstone",
        "   4.1 Session 1: Retrieval-Augmented Generation",
        "   4.2 Session 2: Deployment & Scaling",
        "   4.3 Session 3: Capstone Labs (Track A & B)",
        "   4.4 Session 4: Debrief, Governance & Closing",
        "5. Complete Exercise / Task Reference",
        "6. Summary Statistics",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # ==================== SECTION 1: PROGRAM OVERVIEW ====================
    add_styled_heading(doc, "1. Program Overview & 3-Day Training Flow", level=1)

    p = doc.add_paragraph()
    p.add_run("Program Theme: ").bold = True
    p.add_run(
        "A progressive 3-day journey from LLM fundamentals through production-ready "
        "AI systems, using McKinsey consulting scenarios throughout."
    )

    p = doc.add_paragraph()
    p.add_run("Target Audience: ").bold = True
    p.add_run("Consultants, data scientists, and engineers building AI-powered consulting tools.")

    p = doc.add_paragraph()
    p.add_run("Key Technologies: ").bold = True
    p.add_run("OpenAI API, LangChain, LangGraph, ChromaDB, Pydantic, tiktoken, matplotlib")

    doc.add_paragraph()
    add_styled_heading(doc, "High-Level 3-Day Flow", level=2)

    # Flow table
    flow_table = doc.add_table(rows=1, cols=4)
    flow_table.style = "Table Grid"
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(flow_table, ["", "Day 1", "Day 2", "Day 3"], header=True)
    add_table_row(
        flow_table,
        ["Theme", "Foundations, Prompting\n& Evaluation", "LangChain, LangGraph\n& Multi-Agent", "RAG, Deployment\n& Capstone"],
        bold=True,
        shade_color="E8F0FE",
    )
    add_table_row(
        flow_table,
        [
            "Session 1",
            "Modern LLM Foundations\nfor Agentic Systems",
            "OpenAI API Engineering\nwith Structured Outputs",
            "Retrieval-Augmented\nGeneration (RAG)",
        ],
    )
    add_table_row(
        flow_table,
        [
            "Session 2",
            "Prompt Engineering\nfor Agentic Behaviors",
            "LangChain Development\n& Tool Integration",
            "Deployment &\nScaling",
        ],
    )
    add_table_row(
        flow_table,
        [
            "Session 3",
            "Model Evaluation\n& Comparison",
            "LangGraph Orchestration\n& Workflow Design",
            "Capstone Labs\n(Track A / Track B)",
        ],
    )
    add_table_row(
        flow_table,
        [
            "Session 4",
            "Lab Review &\nIntegration",
            "Multi-Agent Workflows\n& Agentic Systems",
            "Debrief, Governance\n& Closing",
        ],
    )

    # remove initial empty row
    flow_table._tbl.remove(flow_table.rows[0]._tr)

    doc.add_page_break()

    # ==================== DAY 1 ====================
    add_styled_heading(doc, "2. Day 1: Foundations, Prompting & Evaluation", level=1)
    p = doc.add_paragraph()
    p.add_run("Goal: ").bold = True
    p.add_run(
        "Establish core LLM skills -- API usage, token management, prompt engineering, "
        "and model evaluation -- all through McKinsey consulting lenses."
    )

    # --- Session 1 ---
    add_styled_heading(doc, "2.1 Session 1: Modern LLM Foundations for Agentic Systems", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "OpenAI API setup and configuration",
        "Token counting with tiktoken and prompt optimization",
        "Temperature and model parameter experimentation",
        "Conversation history management and multi-turn context",
        "Building a simple conversational agent foundation",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "API Setup & Your First LLM Call", "Initialize OpenAI client, send a consulting prompt, inspect response object"),
        ("Demo 2", "Token Counting & Prompt Budgeting", "Use tiktoken to count tokens, understand pricing implications"),
        ("Demo 3", "Temperature & Sampling Parameters", "Compare outputs at temperatures 0.0, 0.5, 1.0 for the same consulting prompt"),
        ("Demo 4", "Conversation History & Context Windows", "Build multi-turn conversation, observe context retention"),
        ("Demo 5", "Embeddings & Semantic Similarity", "Generate embeddings, compute cosine similarity between consulting concepts"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Configure and Test LLM Connections", "Create test_llm_connection() that initializes an OpenAI client, sends a McKinsey-relevant message, and returns the response. Handle authentication and connection errors."),
        ("Task 2", "Analyze Token Usage & Optimize Prompt Length", "Write optimize_prompt(prompt, max_tokens) that counts tokens using tiktoken and truncates at sentence boundaries if over budget. Returns original/optimized token counts."),
        ("Task 3", "Experiment with Model Parameters", "Create compare_temperatures(prompt) that sends the same prompt at 3 temperature settings (0.0, 0.5, 1.0) and returns comparison dict."),
        ("Task 4", "Build a Simple Conversational Agent", "Build a SimpleAgent class with conversation history, McKinsey engagement manager persona, and multi-turn chat capability. Methods: __init__(), chat(), get_history()."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 2 ---
    add_styled_heading(doc, "2.2 Session 2: Prompt Engineering for Agentic Behaviors", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "System prompt design for specialized agent personas",
        "ReAct (Reason + Act) prompting pattern",
        "Prompt template libraries with variable placeholders",
        "Multi-step reasoning with tool-use descriptions",
        "Structured JSON output from LLMs",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "System Prompt Design Principles", "Craft persona-driven system prompts for a McKinsey research analyst agent"),
        ("Demo 2", "Few-Shot Prompting for Structured Output", "Use examples to guide the model into producing consistent JSON outputs"),
        ("Demo 3", "Chain-of-Thought Reasoning", "Implement step-by-step reasoning for complex consulting analysis"),
        ("Demo 4", "ReAct Pattern Implementation", "Build Thought-Action-Observation loops for analytical reasoning"),
        ("Demo 5", "Prompt Chaining & Decomposition", "Break complex tasks into sequential LLM calls"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Design System Prompts for a Research Agent", "Create a system prompt for a McKinsey Industry Research Agent with MECE breakdowns, data-driven reasoning, and structured output sections (Executive Summary, Key Findings, Strategic Implications, Confidence Level, Data Gaps)."),
        ("Task 2", "Implement ReAct-Style Prompting", "Build a react_agent(question) function using Thought/Action/Observation cycles for M&A due diligence. Available tools: MarketData, FinancialAnalysis, CompetitorLookup."),
        ("Task 3", "Create a Reusable Prompt Template Library", "Build a PromptTemplate class with {variable} placeholders, .format() and .validate() methods. Include 3+ pre-built McKinsey templates (market sizing, executive briefing, competitive analysis)."),
        ("Task 4", "Build a Multi-Step Reasoning Agent", "Build a ToolAgent class with an agentic loop: system prompt describes available tools, LLM outputs JSON to call tools (market_database, financial_model, benchmarking), process results, iterate until final recommendation."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 3 ---
    add_styled_heading(doc, "2.3 Session 3: Model Evaluation and Comparison", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "Custom evaluation rubrics for consulting AI quality",
        "LLM-as-a-Judge evaluation methodology",
        "Multi-model comparison frameworks",
        "Automated evaluation pipelines",
        "Visualization of evaluation results (bar charts, radar charts, scatter plots)",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "Evaluation Rubric Design", "Create scoring criteria tailored to consulting deliverables"),
        ("Demo 2", "LLM-as-a-Judge Implementation", "Use one LLM to score another's output on defined criteria"),
        ("Demo 3", "Automated Test Suite", "Run batch evaluations across multiple prompts and configurations"),
        ("Demo 4", "Results Analysis & Visualization", "Generate charts comparing model performance"),
        ("Demo 5", "Cost-Quality Tradeoff Analysis", "Balance response quality against token costs and latency"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Create Custom Evaluation Rubric", "Build create_agentic_rubric() with 5+ consulting-specific criteria: tool_usage_accuracy, reasoning_quality, mece_decomposition, client_readiness, recommendation_quality. Each scored 1-5."),
        ("Task 2", "Compare Multiple Models on Consulting Tasks", "Build compare_models(prompts, configs, criteria) that runs McKinsey prompts through multiple model configs, scores with LLM-as-Judge, returns a pandas DataFrame with scores, latency, and token counts."),
        ("Task 3", "Build an Automated Evaluation Pipeline", "Build an EvaluationPipeline class: takes (question, expected_behavior) pairs, runs model, evaluates with LLM judge, aggregates scores, identifies weak areas (avg < 3.5), generates summary report."),
        ("Task 4", "Visualize Evaluation Results", "Create plot_criteria_comparison() (bar chart), plot_radar_chart() (spider chart), and plot_latency_vs_quality() (scatter). Include quality threshold line at score 4.0."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 4 ---
    add_styled_heading(doc, "2.4 Session 4: Lab Review & Integration", level=2)
    p = doc.add_paragraph()
    p.add_run("Format: ").bold = True
    p.add_run("Facilitated discussion (no notebook)")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Integration Exercise: ").bold = True
    p.add_run(
        '"Design a consulting AI system on paper" -- participants integrate all Day 1 concepts '
        "(LLM foundations, prompt engineering, model evaluation) into a system design."
    )
    topics = [
        "How would you design a consulting AI system using today's techniques?",
        "Which evaluation criteria matter most for partner-quality consulting deliverables?",
        "How do embeddings and RAG connect to the full RAG pipeline on Day 3?",
        "What are the trade-offs between different prompting strategies?",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_page_break()

    # ==================== DAY 2 ====================
    add_styled_heading(doc, "3. Day 2: LangChain, LangGraph & Multi-Agent Systems", level=1)
    p = doc.add_paragraph()
    p.add_run("Goal: ").bold = True
    p.add_run(
        "Progress from raw API calls to framework-powered agents -- structured outputs, "
        "tool integration, graph-based orchestration, and multi-agent collaboration."
    )

    # --- Session 1 ---
    add_styled_heading(doc, "3.1 Session 1: OpenAI API Engineering with Structured Outputs", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "Streaming responses and token usage tracking",
        "JSON mode for structured data extraction",
        "Function calling with tool schemas",
        "Pydantic-based response validation",
        "Building data extraction pipelines",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "Streaming, Token Usage & Finish Reasons", "Track token usage, stream CEO briefing responses"),
        ("Demo 2", "Structured Output with JSON Mode", "Extract structured client profiles using JSON mode"),
        ("Demo 3", "Function Calling -- Consulting Research Tools", "Define market_research tool schema, execute function-calling loop"),
        ("Demo 4", "Pydantic-Based Response Validation", "Validate LLM output against EngagementSummary Pydantic model"),
        ("Demo 5", "Structured Data Extraction Pipeline", "Reusable pipeline to extract team member info from unstructured emails"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Structured Competitive Assessment Extractor", "Build extract_competitive_intelligence(briefing_text) using JSON mode to extract companies, executives, financial_metrics, strategic_actions, and industries from M&A briefings."),
        ("Task 2", "Financial Analysis Tool with Function Calling", "Create a financial_analysis tool (analysis_type: valuation/profitability/leverage/liquidity, company) using OpenAI function calling. Handle full tool-call loop."),
        ("Task 3", "Multi-Tool Consulting Agent", "Build agent with 3 tools: market_research(industry, region), financial_analysis(company, metric), benchmarking(company, peer_group). Implement dispatch dictionary for automatic routing."),
        ("Task 4", "Robust Consulting API Client", "Build RobustConsultingClient class with: exponential backoff retry (max 3), Pydantic validation via model_validate_json(), streaming support, cumulative token/cost tracking."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 2 ---
    add_styled_heading(doc, "3.2 Session 2: LangChain Development & Tool Integration", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "ChatModels, PromptTemplates, and LCEL (pipe operator)",
        "Building chains: prompt | llm | parser",
        "Custom tool creation with @tool decorator",
        "Document loading and text splitting",
        "Simple RAG pipeline with keyword retrieval",
        "Conversation memory with MessagesPlaceholder",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "LangChain Basics -- ChatModels & PromptTemplates", "Initialize ChatOpenAI, use ChatPromptTemplate with consulting prompts"),
        ("Demo 2", "Building Chains with LCEL", "Build prompt|llm|parser chains, JSON output chains, multi-step chains"),
        ("Demo 3", "Creating Custom Tools", "Build market_sizing_tool, benchmark_lookup_tool, financial_ratio_tool"),
        ("Demo 4", "Document Loading & Text Splitting", "Split McKinsey docs with RecursiveCharacterTextSplitter, observe overlap"),
        ("Demo 5", "Simple RAG Pipeline", "Keyword-based retrieval over McKinsey knowledge base with LCEL chain"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Consulting Analysis Chain", "Create LCEL chain with McKinsey persona, industry/analysis_type inputs, returns structured JSON (key_findings, strategic_recommendation, risk_factors) via JsonOutputParser."),
        ("Task 2", "Custom Consulting Tools", "Create 3 tools: competitor_analysis_tool, market_sizing_estimate (TAM/SAM/SOM), engagement_roi_tool. Bind all to an LLM."),
        ("Task 3", "Conversational Consulting Advisor with Memory", "Build ConsultingAdvisor class using MessagesPlaceholder('history'), maintaining HumanMessage/AIMessage list, McKinsey senior partner persona. Methods: chat(), get_history()."),
        ("Task 4", "RAG-Powered Consulting Knowledge Base", "Build ConsultingRAG class: split Documents with RecursiveCharacterTextSplitter, keyword-overlap scoring retrieval, RAG chain generating grounded MECE answers with source metadata."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 3 ---
    add_styled_heading(doc, "3.3 Session 3: LangGraph Orchestration & Workflow Design", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "LangGraph StateGraph fundamentals",
        "Linear and conditional workflows",
        "ReAct agent pattern with cyclic graphs",
        "Iterative refinement loops (draft/review cycles)",
        "Human-in-the-loop approval gates",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "LangGraph Basics -- Consulting Engagement Pipeline", "StateGraph with 3 nodes: extract_scope, classify_industry, format_brief"),
        ("Demo 2", "Conditional Routing by Complexity", "LLM classifies engagement complexity, routes to rapid/standard/transformation"),
        ("Demo 3", "ReAct Agent -- Market Research Analyst", "Think-Act-Observe cycle with simulated research tools, cyclic graph"),
        ("Demo 4", "Iterative Refinement Cycles", "Draft/refine recommendation until partner-quality (max 3 iterations)"),
        ("Demo 5", "Human-in-the-Loop -- Partner Sign-Off", "Approval gate before client delivery"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run("Student notebooks use generic scenarios; instructor notebooks use McKinsey-themed equivalents.")
    exercises = [
        ("Task 1", "Linear Workflow Pipeline", "Student: 3-node pipeline (clean_node, summarize_node, translate_node). Instructor: Client onboarding pipeline (gather_data, analyze_situation, prepare_brief)."),
        ("Task 2", "Conditional Routing Agent", "Student: Support agent classifying messages (billing/technical/general). Instructor: PE Due Diligence Router by deal size (<$100M, $100M-$1B, >$1B)."),
        ("Task 3", "Self-Correcting Workflow", "Student: Code generation with test/fix loop (max 3 attempts). Instructor: MECE Analysis Refinement Loop -- generate, test MECE criteria, refine until quality passes."),
        ("Task 4", "Research Agent with Planning & Execution", "Student: plan_node (3-step plan), execute_node, check_node, synthesize_node. Instructor: Market Entry Assessment Agent with scoping, workstream execution, and recommendation synthesis."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 4 ---
    add_styled_heading(doc, "3.4 Session 4: Multi-Agent Workflows & Agentic Systems", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "Supervisor-Worker pattern (partner delegates to associates)",
        "Agent-to-agent handoff with context preservation",
        "Parallel agent execution (concurrent workstreams)",
        "Collaborative deliverable creation (team-based document assembly)",
        "End-to-end multi-agent engagement system",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "Supervisor-Worker Pattern", "Partner routes to financial_analyst, strategy_consultant, operations_expert"),
        ("Demo 2", "Agent-to-Agent Handoff", "Three-agent pipeline with structured handoff context passing"),
        ("Demo 3", "Parallel Agent Execution", "Financial, market, competitive agents run concurrently; synthesis merges results"),
        ("Demo 4", "Collaborative Deliverable Creation", "EM storyline -> Associate content -> Senior Partner polish"),
        ("Demo 5", "End-to-End Multi-Agent System", "Triage Partner classifies, routes to specialist, QA check before delivery"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Two-Agent Supervisor-Worker System", "Student: supervisor classifies as 'analyst' or 'creative'. Instructor: Engagement Manager routes to Quantitative or Strategy Associate, with Partner Review."),
        ("Task 2", "Agent Handoff with Context Preservation", "Student: researcher -> analyst -> presenter pipeline with accumulated handoff_context. Instructor: Research Associate -> Industry Expert -> Partner investment thesis pipeline."),
        ("Task 3", "Parallel Research & Synthesis Pipeline", "Student: 3 research agents (technical/business/social) + synthesis. Instructor: Commercial Excellence, Supply Chain, Org Effectiveness workstreams + Senior Partner synthesis."),
        ("Task 4", "Complete Multi-Agent Architecture", "Student: Open-ended -- pick a use case, build 3+ agents, 1+ conditional edge. Instructor: M&A Due Diligence Review System with triage, commercial/operational DD, risk assessment, deal committee report."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    doc.add_page_break()

    # ==================== DAY 3 ====================
    add_styled_heading(doc, "4. Day 3: RAG, Deployment & Capstone", level=1)
    p = doc.add_paragraph()
    p.add_run("Goal: ").bold = True
    p.add_run(
        "Build production-ready RAG systems, learn deployment patterns, and apply everything "
        "in a capstone project with governance considerations."
    )

    # --- Session 1 ---
    add_styled_heading(doc, "4.1 Session 1: Retrieval-Augmented Generation (RAG)", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "Embedding models and cosine similarity search",
        "Vector stores with ChromaDB",
        "Advanced chunking strategies (fixed-size, section-aware, sentence-level)",
        "Query transformation (multi-query expansion, HyDE)",
        "End-to-end RAG pipeline with source citations",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "Embedding Models & Semantic Search", "Create embeddings for consulting snippets, compute similarity, perform semantic search"),
        ("Demo 2", "Vector Stores with ChromaDB", "Create collection, index 8 McKinsey docs with metadata, similarity search"),
        ("Demo 3", "Advanced Chunking Strategies", "Compare fixed-size, section-aware, and sentence-level splitting"),
        ("Demo 4", "Query Transformation Techniques", "Multi-query expansion and HyDE for comprehensive retrieval"),
        ("Demo 5", "End-to-End RAG with Citations", "Build knowledge base, index in ChromaDB, generate answers with source citations"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Embedding-Based Knowledge Search Engine", "Build SearchEngine class: embed documents with OpenAI, search(query, k) returning top-k results ranked by cosine similarity."),
        ("Task 2", "Multi-Strategy Chunking Pipeline", "Build SmartChunker class: detect document type (markdown/code/plain text), apply appropriate splitting strategy, return chunks with metadata."),
        ("Task 3", "Query Expansion & Reranking System", "Build AdvancedRetriever: expand query into 3 formulations, retrieve candidates from ChromaDB, deduplicate, rerank with LLM scoring (0-10)."),
        ("Task 4", "Production RAG with Evaluation Metrics", "Build EvaluatedRAG: query() returns answer + sources + metrics. Evaluate retrieval relevance, answer faithfulness, and answer completeness (all 1-5)."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 2 ---
    add_styled_heading(doc, "4.2 Session 2: Deployment & Scaling", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "API service design with dataclasses and validation",
        "Semantic caching for consulting queries",
        "Monitoring and structured logging",
        "Model routing by query complexity",
        "Token usage and cost tracking",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "Consulting AI API Service Design", "Request/Response dataclasses, McKinseyAIService with multiple service types"),
        ("Demo 2", "Semantic Caching", "Embedding-based similarity cache (threshold 0.90), cache hit/miss demonstration"),
        ("Demo 3", "Monitoring & Structured Logging", "ConsultingAIMonitor with metrics aggregation across services"),
        ("Demo 4", "Model Routing by Complexity", "Classify simple/medium/complex, route to gpt-4o-mini or gpt-4o"),
        ("Demo 5", "Token Usage & Cost Tracking", "Per-model pricing, per-engagement cost tracking, budget alerts"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Rate-Limited LLM Service", "Build RateLimitedLLM: track requests/tokens in 60-second sliding window, block requests exceeding limits, return helpful error messages."),
        ("Task 2", "Response Streaming Simulator", "Build StreamingHandler: use OpenAI streaming API, track time-to-first-token (TTFT) and tokens-per-second, collect full response while reporting progress."),
        ("Task 3", "Multi-Tier Caching System", "Build TieredCache: exact-match tier (O(1) dict), semantic-match tier (embedding similarity), LLM fallback. Report which tier served the response."),
        ("Task 4", "Full Production LLM Gateway", "Build LLMGateway combining: model routing, caching, rate limiting, monitoring, and cost tracking. Includes get_dashboard() for operational metrics."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 3 ---
    add_styled_heading(doc, "4.3 Session 3: Capstone Labs", level=2)
    p = doc.add_paragraph()
    p.add_run("Format: ").bold = True
    p.add_run("Students choose one of two capstone tracks and build incrementally through 4 milestones.")

    # Track A
    add_styled_heading(doc, "Track A: Production RAG -- McKinsey Knowledge Assistant", level=3)
    exercises = [
        ("Milestone 1", "Document Ingestion Pipeline", "Build DocumentIngester: accept McKinsey docs with metadata (practice_area, industry, confidentiality), chunk with RecursiveCharacterTextSplitter, embed with OpenAI, store in ChromaDB with rich metadata."),
        ("Milestone 2", "Advanced Retrieval Engine", "Build RetrievalEngine: expand consulting queries into multiple variants, retrieve from ChromaDB per variant, deduplicate, rerank with LLM-as-judge (0-10 strategic relevance)."),
        ("Milestone 3", "Generation Layer with Citations", "Build RAGGenerator: format context with source tags, generate executive-ready answers with inline citations, flag low-confidence answers when context is insufficient."),
        ("Milestone 4", "Production Wrapper", "Build ProductionRAGService combining Milestones 1-3 + caching, monitoring, and LLM evaluation (strategic_relevance, faithfulness, executive_readiness). Includes dashboard()."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Milestone", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # Track B
    add_styled_heading(doc, "Track B: Multi-Agent -- McKinsey Engagement Team", level=3)
    exercises = [
        ("Milestone 1", "Engagement Manager Agent", "Build EngagementManagerAgent: decompose complex client questions into 2-4 workstreams, assign specialists, return structured engagement plan as JSON."),
        ("Milestone 2", "Specialized Consulting Agents", "Build 4 agent classes: StrategyAnalystAgent, FinancialAnalystAgent, OperationsExpertAgent, IndustryResearcherAgent. Each with McKinsey-style persona and cross-workstream context sharing."),
        ("Milestone 3", "LangGraph Orchestration", "Build StateGraph: EM node creates plan, dispatcher routes to specialist, specialist executes workstream, loop until all complete, synthesis node combines into client deliverable."),
        ("Milestone 4", "Partner Review & QA Loop", "Add EnhancedEngagementState with scores/revision_count. Partner review evaluates analytical_rigor, actionability, strategic_coherence (1-5). Revise weakest area if avg < 3.5, max 2 cycles."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Milestone", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    # --- Session 4 ---
    add_styled_heading(doc, "4.4 Session 4: Debrief, Governance & Closing", level=2)

    p = doc.add_paragraph()
    p.add_run("Topics Covered:").bold = True
    topics = [
        "Cross-track capstone presentations (2 min each)",
        "Safety guardrails (regex + LLM-based intent classification)",
        "Content filtering and output validation",
        "Audit logging for AI-assisted engagements",
        "AI Governance checklist (8 criteria: G1-G8)",
        "Governed agent combining guardrails + filtering + logging",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    add_styled_heading(doc, "Demos (Instructor-Led, Follow-Along)", level=3)
    demos = [
        ("Demo 1", "Safety Guardrails for AI Systems", "Regex pattern blocking + LLM intent classification for consulting scenarios"),
        ("Demo 2", "Content Filtering & Output Validation", "Validate length, relevance, and non-toxicity of consulting outputs"),
        ("Demo 3", "Audit Logging for Engagements", "Structured event logging with user_id, severity, engagement tracking"),
        ("Demo 4", "AI Governance Checklist Evaluator", "8-item checklist (G1-G8) covering data protection, quality, transparency, audit, reliability, cost, bias, privacy"),
        ("Demo 5", "Governed Consulting Agent", "End-to-end governed agent: guardrails + content filter + audit logging"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Demo", "Title", "Description"], header=True)
    for d in demos:
        add_table_row(t, d)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Student Exercises", level=3)
    exercises = [
        ("Task 1", "Input/Output Guardrails", "Build PolicyGuardrail class: load rules from config (pattern, severity: block/warn/allow, description), check input against rules, check output for policy violations."),
        ("Task 2", "Bias Detection Pipeline", "Build BiasDetector: template query with {demographic} placeholder, generate responses for multiple groups, compare with LLM-as-judge for consistency (1-5), report bias with examples."),
        ("Task 3", "Governance Scorecard for Capstone", "Build GovernanceScorecard: evaluate capstone against 8+ criteria (Client Data Protection, Output Quality, Bias Testing, Audit Trail, Transparency, Reliability, Human Oversight, Financial Validation). Score 1-5 with justification, identify top 3 gaps, propose remediations."),
        ("Task 4", "Deployment Readiness Assessment", "Build DeploymentReadiness: check technical (error handling, caching, monitoring), governance (guardrails, audit, bias testing), and operational readiness (documentation, runbooks, alerts). Produce go/no-go recommendation with blocking issues."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_row(t, ["Task", "Title", "Description"], header=True)
    for e in exercises:
        add_table_row(t, e)
    t._tbl.remove(t.rows[0]._tr)

    add_styled_heading(doc, "Closing Reflection Questions", level=3)
    questions = [
        "How should McKinsey govern AI use in client engagements? What approval workflows are needed?",
        "What guardrails are essential for AI-generated strategy recommendations?",
        "How do we maintain client trust when AI contributes to deliverables?",
        "What bias testing should be standard before deploying AI for industry/geographic recommendations?",
    ]
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. {q}")

    doc.add_page_break()

    # ==================== SECTION 5: COMPLETE EXERCISE REFERENCE ====================
    add_styled_heading(doc, "5. Complete Exercise / Task Reference", level=1)
    p = doc.add_paragraph(
        "Master table of all hands-on exercises, tasks, and capstone milestones across the 3-day program."
    )

    ref_table = doc.add_table(rows=1, cols=5)
    ref_table.style = "Table Grid"
    add_table_row(ref_table, ["Day", "Session", "Exercise", "Type", "Student TODO?"], header=True)

    # Day 1
    d1_exercises = [
        ["1", "S1", "Task 1: Configure & Test LLM Connections", "Code", "Yes"],
        ["1", "S1", "Task 2: Token Usage & Prompt Optimization", "Code", "Yes"],
        ["1", "S1", "Task 3: Model Parameter Experimentation", "Code", "Yes"],
        ["1", "S1", "Task 4: Simple Conversational Agent", "Code", "Yes"],
        ["1", "S2", "Task 1: System Prompts for Research Agent", "Code", "Yes"],
        ["1", "S2", "Task 2: ReAct-Style Prompting", "Code", "Yes"],
        ["1", "S2", "Task 3: Prompt Template Library", "Code", "Yes"],
        ["1", "S2", "Task 4: Multi-Step Reasoning Agent", "Code", "Yes"],
        ["1", "S3", "Task 1: Custom Evaluation Rubric", "Code", "Yes"],
        ["1", "S3", "Task 2: Multi-Model Comparison", "Code", "Yes"],
        ["1", "S3", "Task 3: Automated Evaluation Pipeline", "Code", "Yes"],
        ["1", "S3", "Task 4: Evaluation Visualization", "Code", "Yes"],
        ["1", "S4", "Integration Exercise: System Design", "Discussion", "N/A"],
    ]
    for row in d1_exercises:
        add_table_row(ref_table, row)

    # Day 2
    d2_exercises = [
        ["2", "S1", "Task 1: Structured Competitive Extractor", "Code", "Yes"],
        ["2", "S1", "Task 2: Financial Analysis Function Calling", "Code", "Yes"],
        ["2", "S1", "Task 3: Multi-Tool Consulting Agent", "Code", "Yes"],
        ["2", "S1", "Task 4: Robust API Client", "Code", "Yes"],
        ["2", "S2", "Task 1: Consulting Analysis Chain (LCEL)", "Code", "Yes"],
        ["2", "S2", "Task 2: Custom Consulting Tools", "Code", "Yes"],
        ["2", "S2", "Task 3: Conversational Advisor with Memory", "Code", "Yes"],
        ["2", "S2", "Task 4: RAG-Powered Knowledge Base", "Code", "Yes"],
        ["2", "S3", "Task 1: Linear Workflow Pipeline", "Code", "Yes"],
        ["2", "S3", "Task 2: Conditional Routing Agent", "Code", "Yes"],
        ["2", "S3", "Task 3: Self-Correcting Workflow", "Code", "Yes"],
        ["2", "S3", "Task 4: Research Agent with Planning", "Code", "Yes"],
        ["2", "S4", "Task 1: Supervisor-Worker System", "Code", "Yes"],
        ["2", "S4", "Task 2: Agent Handoff with Context", "Code", "Yes"],
        ["2", "S4", "Task 3: Parallel Research & Synthesis", "Code", "Yes"],
        ["2", "S4", "Task 4: Complete Multi-Agent Architecture", "Code", "Yes"],
    ]
    for row in d2_exercises:
        add_table_row(ref_table, row)

    # Day 3
    d3_exercises = [
        ["3", "S1", "Task 1: Embedding-Based Search Engine", "Code", "Yes"],
        ["3", "S1", "Task 2: Multi-Strategy Chunking Pipeline", "Code", "Yes"],
        ["3", "S1", "Task 3: Query Expansion & Reranking", "Code", "Yes"],
        ["3", "S1", "Task 4: Production RAG with Evaluation", "Code", "Yes"],
        ["3", "S2", "Task 1: Rate-Limited LLM Service", "Code", "Yes"],
        ["3", "S2", "Task 2: Response Streaming Simulator", "Code", "Yes"],
        ["3", "S2", "Task 3: Multi-Tier Caching System", "Code", "Yes"],
        ["3", "S2", "Task 4: Full Production LLM Gateway", "Code", "Yes"],
        ["3", "S3A", "Milestone 1: Document Ingestion Pipeline", "Capstone", "Yes"],
        ["3", "S3A", "Milestone 2: Advanced Retrieval Engine", "Capstone", "Yes"],
        ["3", "S3A", "Milestone 3: Generation with Citations", "Capstone", "Yes"],
        ["3", "S3A", "Milestone 4: Production Wrapper", "Capstone", "Yes"],
        ["3", "S3B", "Milestone 1: Engagement Manager Agent", "Capstone", "Yes"],
        ["3", "S3B", "Milestone 2: Specialized Consulting Agents", "Capstone", "Yes"],
        ["3", "S3B", "Milestone 3: LangGraph Orchestration", "Capstone", "Yes"],
        ["3", "S3B", "Milestone 4: Partner Review & QA Loop", "Capstone", "Yes"],
        ["3", "S4", "Task 1: Input/Output Guardrails", "Code", "Yes"],
        ["3", "S4", "Task 2: Bias Detection Pipeline", "Code", "Yes"],
        ["3", "S4", "Task 3: Governance Scorecard", "Code", "Yes"],
        ["3", "S4", "Task 4: Deployment Readiness Assessment", "Code", "Yes"],
        ["3", "S4", "Cross-Track Capstone Presentations", "Presentation", "N/A"],
    ]
    for row in d3_exercises:
        add_table_row(ref_table, row)

    ref_table._tbl.remove(ref_table.rows[0]._tr)

    doc.add_page_break()

    # ==================== SECTION 6: SUMMARY STATISTICS ====================
    add_styled_heading(doc, "6. Summary Statistics", level=1)

    stats_table = doc.add_table(rows=1, cols=5)
    stats_table.style = "Table Grid"
    add_table_row(stats_table, ["Metric", "Day 1", "Day 2", "Day 3", "Total"], header=True)
    stats = [
        ["Sessions", "4", "4", "4", "12"],
        ["Instructor-Led Demos", "15*", "20", "15", "50"],
        ["Student Coding Tasks", "12", "16", "12", "40"],
        ["Capstone Milestones", "0", "0", "8", "8"],
        ["Discussion/Presentation Activities", "1", "0", "2", "3"],
        ["Total Hands-On Activities", "13", "16", "22", "51"],
    ]
    for s in stats:
        add_table_row(stats_table, s)
    stats_table._tbl.remove(stats_table.rows[0]._tr)

    p = doc.add_paragraph()
    p.add_run("\n* Day 1 Session 3 demos are counted from the student notebook (no separate instructor notebook for Session 3).")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    add_styled_heading(doc, "Technology Progression", level=2)

    prog_table = doc.add_table(rows=1, cols=3)
    prog_table.style = "Table Grid"
    add_table_row(prog_table, ["Day", "Core Technologies", "Key Skills Built"], header=True)
    prog_data = [
        ["Day 1", "OpenAI API, tiktoken, matplotlib", "API calls, prompt engineering, evaluation rubrics, LLM-as-Judge"],
        ["Day 2", "OpenAI Function Calling, LangChain, LangGraph", "Structured outputs, LCEL chains, tools, StateGraph, multi-agent orchestration"],
        ["Day 3", "ChromaDB, OpenAI Embeddings, Pydantic", "RAG pipelines, deployment patterns, caching, governance, production readiness"],
    ]
    for row in prog_data:
        add_table_row(prog_table, row)
    prog_table._tbl.remove(prog_table.rows[0]._tr)

    doc.add_paragraph()
    add_styled_heading(doc, "Notebook Inventory", level=2)
    p = doc.add_paragraph()
    p.add_run("Instructor notebooks: ").bold = True
    p.add_run("11 (Session 1-1 through 3-4, except Day 1 Session 3)")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Student notebooks: ").bold = True
    p.add_run("12 (all sessions including both capstone tracks)")

    # Save
    output_path = "/Users/siddharthsharma/Desktop/mckinsey-genAi-3Day/GenAI_3Day_Training_Guide.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
